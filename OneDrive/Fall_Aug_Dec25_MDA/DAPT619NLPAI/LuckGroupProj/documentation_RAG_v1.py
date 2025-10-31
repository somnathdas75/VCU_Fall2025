# documentation_RAG_v1.py
"""
Documentation Generator for VCU MDA Weekend AI Chatbot
Creates a comprehensive DOCX documentation file with Table of Contents
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os

def create_chatbot_documentation():
    """Create the main documentation file"""
    # Create a new Document
    doc = Document()
    
    # Title Page
    title = doc.add_heading('VCU MDA Weekend AI Chatbot Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Document Version: 2.0")
    doc.add_paragraph(f"Created Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph("Author: AI Assistant")
    doc.add_page_break()

    # Add placeholder for Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.add_run("Table of Contents will be generated automatically...").italic = True
    doc.add_page_break()

    # Rest of your documentation content remains the same...
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_heading('1.1 Overview', level=2)
    doc.add_paragraph(
        "The VCU MDA Weekend AI Chatbot is an intelligent virtual assistant designed to provide "
        "prospective students with comprehensive information about the Master of Decision Analytics "
        "Weekend program at Virginia Commonwealth University. The system leverages advanced Natural "
        "Language Processing (NLP) and Artificial Intelligence (AI) technologies to deliver accurate, "
        "context-aware responses in real-time."
    )
    
    doc.add_heading('1.2 Key Features', level=2)
    features = [
        "AI-Powered Natural Language Understanding",
        "Semantic Search with Sentence Transformers",
        "Real-time Content Fetching from VCU Sources",
        "Intelligent Response Generation using OpenAI GPT",
        "Conversation Memory with Chroma DB",
        "Multi-source Information Integration",
        "User Intent Recognition and Classification",
        "Advanced Entity Extraction"
    ]
    
    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature).bold = True

    doc.add_heading('1.3 Technology Stack', level=2)
    tech_stack = [
        ("Frontend", "Streamlit - Python web framework for interactive applications"),
        ("NLP Processing", "spaCy - Industrial-strength natural language processing"),
        ("AI Models", "OpenAI GPT, Sentence Transformers - Intelligent response generation"),
        ("Vector Database", "Chroma DB - Persistent conversation storage and retrieval"),
        ("Web Scraping", "BeautifulSoup4 - Content extraction from VCU websites"),
        ("Machine Learning", "scikit-learn - Similarity calculations and clustering")
    ]
    
    for tech, description in tech_stack:
        p = doc.add_paragraph()
        p.add_run(f"{tech}: ").bold = True
        p.add_run(description)

    # 2. System Architecture
    doc.add_heading('2. System Architecture', level=1)
    
    doc.add_heading('2.1 High-Level Architecture', level=2)
    doc.add_paragraph(
        "The chatbot follows a modular microservices architecture with the following components:"
    )
    
    architecture_components = [
        ("User Interface Layer", "Streamlit-based web interface handling user interactions"),
        ("NLP Processing Layer", "Intent recognition, entity extraction, and semantic analysis"),
        ("AI Response Layer", "OpenAI integration and rule-based response generation"),
        ("Data Layer", "Chroma DB for conversation storage and vector embeddings"),
        ("Content Layer", "Real-time data fetching from VCU websites and bulletin")
    ]
    
    for component, description in architecture_components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(component).bold = True
        p.add_run(f" - {description}")

    doc.add_heading('2.2 Data Flow', level=2)
    data_flow = [
        "1. User submits query through Streamlit interface",
        "2. Query processed by NLP layer for intent and entity extraction",
        "3. System searches Chroma DB for similar past conversations",
        "4. Real-time content fetched from VCU sources if needed",
        "5. AI response generator creates context-aware answer",
        "6. Response displayed to user and stored in conversation history",
        "7. Conversation metadata saved to Chroma DB for future reference"
    ]
    
    for step in data_flow:
        doc.add_paragraph(step, style='List Number')

    # 3. Installation Guide
    doc.add_heading('3. Installation Guide', level=1)
    
    doc.add_heading('3.1 Prerequisites', level=2)
    prerequisites = [
        "Python 3.8 or higher",
        "pip package manager",
        "Stable internet connection",
        "VCU website access (for content fetching)",
        "OpenAI API key (optional, for enhanced AI features)"
    ]
    
    for prereq in prerequisites:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('3.2 Installation Steps', level=2)
    
    doc.add_heading('Step 1: Clone or Download the Code', level=3)
    doc.add_paragraph("Save the chatbot code as VCU_MDA_V7_AI.py in your preferred directory.")

    doc.add_heading('Step 2: Create Requirements File', level=3)
    doc.add_paragraph("Create a requirements.txt file with the following content:")
    
    requirements_code = """streamlit>=1.28.0
requests>=2.31.0
beautifulsoup4>=4.12.0
pandas>=2.0.0
numpy>=1.24.0
spacy>=3.7.0
sentence-transformers>=2.2.0
scikit-learn>=1.3.0
chromadb>=0.4.0
python-docx>=1.1.0"""
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(requirements_code)
    code_run.font.name = 'Courier New'

    doc.add_heading('Step 3: Install Dependencies', level=3)
    doc.add_paragraph("Open terminal/command prompt and run:")
    cmd_para = doc.add_paragraph()
    cmd_run = cmd_para.add_run("pip install -r requirements.txt")
    cmd_run.font.name = 'Courier New'
    
    doc.add_heading('Step 4: Download spaCy Model', level=3)
    doc.add_paragraph("Install the English language model:")
    spacy_para = doc.add_paragraph()
    spacy_run = spacy_para.add_run("python -m spacy download en_core_web_sm")
    spacy_run.font.name = 'Courier New'

    doc.add_heading('Step 5: Configure API Keys (Optional)', level=3)
    doc.add_paragraph("Create a .streamlit/secrets.toml file for OpenAI API key:")
    
    secrets_code = """[openai]
api_key = "your-openai-api-key-here\""""
    
    secrets_para = doc.add_paragraph()
    secrets_run = secrets_para.add_run(secrets_code)
    secrets_run.font.name = 'Courier New'

    # 4. Configuration
    doc.add_heading('4. Configuration', level=1)
    
    doc.add_heading('4.1 Primary URLs', level=2)
    urls_config = [
        ("PRIMARY_BASE", "https://business.vcu.edu/graduate-programs/mda-weekend/", "Main program page"),
        ("BULLETIN_URL", "https://bulletin.vcu.edu/graduate/school-business/decision-analytics-programs/decision-analytics-mda-pro/", "Official curriculum and requirements"),
        ("APPLY_URL", "https://gradadmissions.vcu.edu/portal/apply?_gl=1*qxhpby*_gcl_au*NzE4MTA2NDk1LjE3NjE3ODc2NTE.*_ga*Mjk1NjQ2NzY5LjE3NTk3NjQ3MTg.*_ga_WMHV0FXMBD*czE3NjE4NjAxNDEkbzEkZzEkdDE3NjE4NjAzMzAkajYwJGwwJGhw", "Application portal"),
        ("REFERRAL_URL", "https://business.vcu.edu/graduate-programs/mda-weekend/referral-award/", "Alumni referral program")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'URL'
    hdr_cells[2].text = 'Description'
    
    for var_name, url, desc in urls_config:
        row_cells = table.add_row().cells
        row_cells[0].text = var_name
        row_cells[1].text = url
        row_cells[2].text = desc

    doc.add_heading('4.2 AI Model Configuration', level=2)
    ai_config = [
        ("Sentence Transformer Model", "all-MiniLM-L6-v2", "Semantic similarity calculations"),
        ("OpenAI Model", "gpt-3.5-turbo", "Intelligent response generation"),
        ("spaCy Model", "en_core_web_sm", "NLP processing and entity recognition"),
        ("Similarity Threshold", "0.3", "Minimum similarity score for context matching"),
        ("Response Temperature", "0.3", "AI creativity control (lower = more factual)")
    ]
    
    for setting, value, purpose in ai_config:
        p = doc.add_paragraph()
        p.add_run(f"{setting}: ").bold = True
        p.add_run(f"{value} - {purpose}")

    # 5. User Guide
    doc.add_heading('5. User Guide', level=1)
    
    doc.add_heading('5.1 Getting Started', level=2)
    doc.add_paragraph("To launch the chatbot:")
    startup_steps = [
        "Open terminal/command prompt in the project directory",
        "Run the command: streamlit run VCU_MDA_V7_AI.py",
        "The application will open in your default web browser",
        "Wait for the AI components to initialize (indicated by status messages)"
    ]
    
    for step in startup_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('5.2 Using the Chatbot', level=2)
    
    doc.add_heading('Basic Operations', level=3)
    basic_ops = [
        ("Ask Questions", "Type natural language questions in the chat input box"),
        ("Quick Actions", "Use sidebar buttons for common questions"),
        ("Load Content", "Click 'Load AI Context' to fetch latest program information"),
        ("View Analysis", "Enable 'Show Analysis' to see NLP processing details")
    ]
    
    for op, desc in basic_ops:
        p = doc.add_paragraph()
        p.add_run(f"{op}: ").bold = True
        p.add_run(desc)

    doc.add_heading('5.3 Example Questions', level=2)
    example_categories = [
        ("Admission Questions", [
            "What are the application deadlines?",
            "What qualifications do I need?",
            "How do I apply to the program?",
            "Is GMAT required for admission?"
        ]),
        ("Program Questions", [
            "How long is the program?",
            "What is the program format?",
            "What will I learn in the curriculum?",
            "Tell me about the practicum experience"
        ]),
        ("Financial Questions", [
            "What is the total cost?",
            "Are there scholarships available?",
            "Tell me about the alumni referral award",
            "What payment options are available?"
        ]),
        ("Advanced AI Questions", [
            "How does this program compare to other analytics programs?",
            "What career outcomes can I expect?",
            "How does the weekend format benefit working professionals?",
            "What makes this program unique?"
        ])
    ]
    
    for category, questions in example_categories:
        doc.add_heading(category, level=3)
        for question in questions:
            doc.add_paragraph(question, style='List Bullet')

    # 6. Technical Details
    doc.add_heading('6. Technical Details', level=1)
    
    doc.add_heading('6.1 Core Classes', level=2)
    
    classes_info = [
        ("AINLPProcessor", """Responsibilities:
- Natural Language Processing operations
- Intent recognition and classification
- Entity extraction using spaCy
- Semantic similarity calculations

Key Methods:
- extract_entities(text): Extract named entities from text
- analyze_intent(query): Determine user intent from query
- semantic_similarity(text1, text2): Calculate similarity score"""),
        
        ("AIResponseGenerator", """Responsibilities:
- Generate intelligent responses using AI
- Manage knowledge base and context
- Handle rule-based and AI-generated responses

Key Methods:
- generate_ai_response(query, context): Main response generation
- find_most_relevant_context(query, chunks): Semantic context matching
- generate_rule_based_response(): Fallback response system"""),
        
        ("AIChromaDBChatHistory", """Responsibilities:
- Persistent conversation storage
- Semantic search across chat history
- Conversation metadata management

Key Methods:
- add_message(): Store conversation with metadata
- get_similar_conversations(): Find related past conversations
- initialize_db(): Chroma DB setup and configuration""")
    ]
    
    for class_name, description in classes_info:
        doc.add_heading(class_name, level=3)
        doc.add_paragraph(description)

    doc.add_heading('6.2 Data Structures', level=2)
    data_structures = [
        ("Knowledge Base", "Hierarchical dictionary organizing program information by categories"),
        ("Conversation Metadata", "JSON structure storing query, response, intent, timestamp, and sources"),
        ("Content Chunks", "List of processed sentences from VCU websites for context"),
        ("Entity Objects", "Structured entities with text, label, and position information")
    ]
    
    for ds, desc in data_structures:
        p = doc.add_paragraph()
        p.add_run(f"{ds}: ").bold = True
        p.add_run(desc)

    # 7. API Reference
    doc.add_heading('7. API Reference', level=1)
    
    doc.add_heading('7.1 Method Specifications', level=2)
    
    methods = [
        ("AINLPProcessor.analyze_intent(query)", """Parameters:
- query (str): User's input question

Returns:
- str: Detected intent category ('admission', 'curriculum', 'financial', 'deadline', 'general_info')

Description:
Analyzes user query to determine primary intent using keyword matching and pattern recognition."""),
        
        ("AIResponseGenerator.generate_ai_response(query, context_chunks)", """Parameters:
- query (str): User's input question
- context_chunks (list): List of relevant content sentences

Returns:
- str: AI-generated or rule-based response

Description:
Main response generation method that combines AI capabilities with contextual information."""),
        
        ("AIChromaDBChatHistory.get_similar_conversations(query, limit=3)", """Parameters:
- query (str): Current user question
- limit (int): Maximum number of similar conversations to return

Returns:
- list: Similar past conversations with similarity scores

Description:
Finds semantically similar past conversations using vector similarity search.""")
    ]
    
    for method, details in methods:
        doc.add_heading(method, level=3)
        doc.add_paragraph(details)

    # 8. Troubleshooting
    doc.add_heading('8. Troubleshooting', level=1)
    
    issues_solutions = [
        ("Chatbot not starting", [
            "Check Python version (requires 3.8+)",
            "Verify all dependencies are installed",
            "Ensure no port conflicts (Streamlit uses port 8501)"
        ]),
        ("AI features not working", [
            "Verify spaCy model is downloaded",
            "Check internet connection for API calls",
            "Validate OpenAI API key in secrets.toml",
            "Check console for import errors"
        ]),
        ("Content fetching failures", [
            "Verify VCU websites are accessible",
            "Check network connectivity",
            "Review error messages in console",
            "Try loading content manually using sidebar button"
        ]),
        ("Slow response times", [
            "Reduce number of simultaneous context chunks",
            "Disable semantic search if not needed",
            "Use rule-based responses instead of AI",
            "Check system resources and memory"
        ])
    ]
    
    for issue, solutions in issues_solutions:
        doc.add_heading(issue, level=2)
        for solution in solutions:
            doc.add_paragraph(solution, style='List Bullet')

    # 9. Maintenance
    doc.add_heading('9. Maintenance', level=1)
    
    doc.add_heading('9.1 Regular Maintenance Tasks', level=2)
    maintenance_tasks = [
        "Update dependencies periodically",
        "Verify VCU URLs are still valid",
        "Refresh content context regularly",
        "Monitor AI API usage and costs",
        "Backup Chroma DB conversation history",
        "Review and update response templates"
    ]
    
    for task in maintenance_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('9.2 Update Procedures', level=2)
    update_steps = [
        "1. Pull latest code changes",
        "2. Update requirements.txt with new dependencies",
        "3. Run dependency installation: pip install -r requirements.txt",
        "4. Test all chatbot features",
        "5. Update documentation if needed",
        "6. Deploy updated version"
    ]
    
    for step in update_steps:
        doc.add_paragraph(step, style='List Number')

    # Save the document
    filename = f"VCU_MDA_Chatbot_Documentation_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    print(f"Documentation saved as: {filename}")
    
    return filename

def create_simple_toc_documentation():
    """Create documentation with a simple manual Table of Contents"""
    doc = Document()
    
    # Title Page
    title = doc.add_heading('VCU MDA Weekend AI Chatbot Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Created: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    doc.add_page_break()

    # Manual Table of Contents
    doc.add_heading('Table of Contents', level=1)
    
    sections = [
        ("1. Executive Summary", "High-level overview and features"),
        ("2. System Architecture", "Technical design and components"),
        ("3. Installation Guide", "Setup and configuration steps"),
        ("4. User Guide", "How to use the chatbot"),
        ("5. Technical Details", "Classes and data structures"),
        ("6. API Reference", "Method specifications"),
        ("7. Troubleshooting", "Common issues and solutions"),
        ("8. Maintenance", "Regular upkeep procedures")
    ]
    
    for section, description in sections:
        p = doc.add_paragraph()
        p.add_run(section).bold = True
        p.add_run(f" - {description}")
    
    doc.add_page_break()
    
    # Add content sections
    for section, description in sections:
        doc.add_heading(section, level=1)
        doc.add_paragraph(f"This section covers {description.lower()}")
        # Add more content here as needed
    
    filename = f"VCU_MDA_Documentation_Simple_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    print(f"Simple documentation with TOC saved as: {filename}")
    return filename

if __name__ == "__main__":
    print("Creating VCU MDA Chatbot Documentation...")
    
    # Option 1: Create full documentation (without Aspose Words TOC)
    doc_file = create_chatbot_documentation()
    print(f"✓ Main documentation created: {doc_file}")
    
    # Option 2: Create simple documentation with manual TOC
    simple_doc = create_simple_toc_documentation()
    print(f"✓ Simple documentation created: {simple_doc}")
    
    print("\n📚 Documentation files created successfully!")
    print("Note: For automatic Table of Contents, you would need:")
    print("1. Microsoft Word (to update TOC fields manually)")
    print("2. Or install Aspose.Words and ensure the file exists")
    print("\nThe current files include manual table of contents.")